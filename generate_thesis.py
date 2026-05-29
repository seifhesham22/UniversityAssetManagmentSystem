"""Generate UAMS Bachelor Thesis .docx (NR TSU format)."""
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup: A4, margins L30 R15 T20 B20 ────────────────────────────────
section = doc.sections[0]
section.page_width  = Mm(210)
section.page_height = Mm(297)
section.left_margin   = Mm(30)
section.right_margin  = Mm(15)
section.top_margin    = Mm(20)
section.bottom_margin = Mm(20)

# ── Helper: apply paragraph formatting ─────────────────────────────────────
def fmt(para, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=0, space_after=6, first_indent_cm=None):
    pf = para.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    if first_indent_cm is not None:
        pf.first_line_indent = Cm(first_indent_cm)
    for run in para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = bold

def add_para(text, size=12, bold=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=6,
             first_indent=1.25):
    p = doc.add_paragraph(text)
    fmt(p, size=size, bold=bold, align=align,
        space_before=space_before, space_after=space_after,
        first_indent_cm=first_indent)
    return p

def add_heading(text, level=1):
    """Chapter / section headings — bold, centered for level 1; left for 2/3."""
    size = 14 if level == 1 else 12
    align = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    sb = 12 if level == 1 else 6
    p = doc.add_paragraph()
    run = p.add_run(text.upper() if level == 1 else text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(sb)
    pf.space_after  = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0)
    return p

def add_page_break():
    doc.add_page_break()

def body(text, first_indent=1.25):
    return add_para(text, first_indent=first_indent)

# ═══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
pf = p.paragraph_format
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf.space_before = Pt(0); pf.space_after = Pt(0)
run = p.add_run("MINISTRY OF SCIENCE AND HIGHER EDUCATION OF THE RUSSIAN FEDERATION\n"
                "Federal State Autonomous Educational Institution of Higher Education\n"
                "\"National Research Tomsk State University\"\n"
                "Faculty of Innovative Technologies\n"
                "Department of Software Engineering")
run.font.name = "Times New Roman"; run.font.size = Pt(12)

doc.add_paragraph()

p = add_para("Seif Hesham Elmoazen", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0)

doc.add_paragraph()

p = add_para(
    "UNIVERSITY ASSET MANAGEMENT SYSTEM (UAMS):\n"
    "A WEB-BASED PLATFORM FOR MANAGING PHYSICAL ASSETS IN UNIVERSITY FACILITIES",
    bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0)

doc.add_paragraph()

p = add_para("Bachelor's Thesis", size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0)
p.runs[0].italic = True

doc.add_paragraph()
doc.add_paragraph()

p = add_para("Specialization: 09.03.04 Software Engineering", size=12,
             align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=0)

doc.add_paragraph()

p = add_para("Academic Supervisor: _______________  [Supervisor Name], [Title]",
             size=12, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=0)
p2 = add_para("Student: _______________  Seif Hesham Elmoazen",
              size=12, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=0)

doc.add_paragraph()
doc.add_paragraph()

p = add_para("Tomsk 2026", size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════
add_heading("ABSTRACT", level=1)

body(
    "This bachelor's thesis presents the design and implementation of the University Asset "
    "Management System (UAMS) — a full-stack web application developed to replace paper-based "
    "and spreadsheet-driven processes for tracking physical assets inside university buildings. "
    "The system enables asset managers to create interactive, drag-and-drop room layouts on a "
    "canvas, place asset definitions, record asset conditions, manage maintenance checklists, "
    "and receive defect reports through a dedicated Telegram/VK social-media bot integration."
)
body(
    "The backend is built with ASP.NET Core following Clean Architecture principles, using the "
    "CQRS pattern with MediatR and Entity Framework Core with a JSON-owned column strategy for "
    "storing canvas layout data. The frontend is a React/TypeScript single-page application "
    "with a custom canvas renderer that supports zoom, pan, grouping, and asset linking. "
    "The work demonstrates how modern web technologies can be combined with domain-driven "
    "design to produce a maintainable, role-aware enterprise application."
)
body("Keywords: asset management, university facilities, web application, React, ASP.NET Core, "
     "CQRS, canvas designer, checklist, VK bot.")

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual — matches actual content)
# ═══════════════════════════════════════════════════════════════════════════
add_heading("CONTENTS", level=1)

toc_entries = [
    ("Glossary", "4"),
    ("Introduction", "5"),
    ("1  Design", "7"),
    ("   1.1  Functional Requirements", "7"),
    ("   1.2  Problem Domain Model", "9"),
    ("2  Technologies and Tools", "11"),
    ("   2.1  React, TypeScript and Vite", "11"),
    ("   2.2  ASP.NET Core", "12"),
    ("      2.2.1  Entity Framework Core", "13"),
    ("      2.2.2  MediatR and CQRS", "14"),
    ("   2.3  SQL Server", "15"),
    ("3  Architecture", "16"),
    ("   3.1  API Layer", "16"),
    ("   3.2  Application Layer", "17"),
    ("   3.3  Persistence Layer", "18"),
    ("   3.4  Domain Layer", "18"),
    ("4  Implementation", "19"),
    ("   4.1  Authentication and Role-Based Access Control", "19"),
    ("   4.2  Asset Definitions", "20"),
    ("   4.3  Buildings and Rooms", "21"),
    ("   4.4  Canvas Designer", "22"),
    ("      4.4.1  Placing, Moving and Resizing", "22"),
    ("      4.4.2  Grouping and Room Linking", "24"),
    ("   4.5  Asset Conditions and Checklists", "25"),
    ("   4.6  Reporting and VK Bot Integration", "27"),
    ("Conclusion", "29"),
    ("References", "30"),
    ("Appendix A – Selected Screenshots", "31"),
]

for title, page in toc_entries:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(0); pf.space_after = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0)
    tab_stops = pf.tab_stops
    # add a right-aligned tab at right margin position
    from docx.shared import Inches
    run = p.add_run(f"{title}")
    run.font.name = "Times New Roman"; run.font.size = Pt(12)
    # add dots + page number
    run2 = p.add_run(f" {'.' * max(1, 60 - len(title) - len(page))} {page}")
    run2.font.name = "Times New Roman"; run2.font.size = Pt(12)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# GLOSSARY
# ═══════════════════════════════════════════════════════════════════════════
add_heading("GLOSSARY", level=1)

glossary = [
    ("API",
     "Application Programming Interface — a defined contract through which one software "
     "component exposes functionality to another, typically over HTTP in this project."),
    ("ASP.NET Core",
     "An open-source, cross-platform web framework by Microsoft used to build the UAMS "
     "backend REST API."),
    ("Asset",
     "Any physical item — furniture, equipment, or infrastructure — tracked inside a "
     "university facility by the UAMS."),
    ("Asset Definition",
     "A reusable template describing a category of asset (name, SVG icon, checklist "
     "template). Multiple placed instances can share one definition."),
    ("Canvas",
     "The interactive 2-D drawing surface in the frontend where asset managers arrange "
     "placed assets inside a room layout."),
    ("Checklist",
     "A list of maintenance or inspection tasks associated with a placed asset for a "
     "specific study year."),
    ("Clean Architecture",
     "A layered software design approach (Domain → Application → Infrastructure → "
     "Presentation) that keeps business rules independent of frameworks and databases."),
    ("CQRS",
     "Command Query Responsibility Segregation — a pattern that separates write operations "
     "(commands) from read operations (queries) to improve clarity and scalability."),
    ("CSS",
     "Cascading Style Sheets — the language used to describe the visual presentation of "
     "HTML documents."),
    ("DTO",
     "Data Transfer Object — a simple object used to carry data between layers or across "
     "a network boundary without exposing domain internals."),
    ("Entity Framework Core",
     "A .NET object-relational mapper (ORM) that lets developers work with a SQL Server "
     "database using C# objects instead of raw SQL."),
    ("Faculty",
     "The top-level organisational unit in UAMS (e.g., Faculty of Innovative Technologies) "
     "that owns buildings and their assets."),
    ("GUID",
     "Globally Unique Identifier — a 128-bit number used as a primary key throughout the "
     "UAMS database to avoid collisions and enable distributed generation."),
    ("HTML",
     "HyperText Markup Language — the standard markup language for structuring web pages."),
    ("HTTP",
     "HyperText Transfer Protocol — the communication protocol used for all API calls "
     "between the UAMS frontend and backend."),
    ("JWT",
     "JSON Web Token — a compact, self-contained token format used for authentication "
     "and transmitting role claims between the identity service and UAMS API."),
    ("Layout",
     "The persisted snapshot of all placed assets inside a room, stored as a JSON-owned "
     "column in the Room entity."),
    ("MediatR",
     "A .NET library implementing the Mediator pattern; used in UAMS to dispatch CQRS "
     "commands and queries to their respective handlers."),
    ("ORM",
     "Object-Relational Mapper — software that converts between in-memory objects and "
     "relational database rows automatically."),
    ("Placed Asset",
     "A concrete instance of an Asset Definition positioned at specific coordinates inside "
     "a room layout on the canvas."),
    ("REST",
     "Representational State Transfer — an architectural style for web services that uses "
     "standard HTTP methods and stateless communication."),
    ("Role-Based Access Control (RBAC)",
     "A security model that restricts system operations to users who hold the appropriate "
     "role (Asset Manager or Reporter in UAMS)."),
    ("SPA",
     "Single-Page Application — a web application that loads a single HTML page and "
     "dynamically updates content in the browser without full page reloads."),
    ("SQL Server",
     "Microsoft's relational database management system used as the UAMS data store."),
    ("SVG",
     "Scalable Vector Graphics — an XML-based vector image format used for asset icons "
     "on the canvas because it scales without loss of quality."),
    ("TSX",
     "TypeScript XML — a file extension for TypeScript files that contain JSX syntax, "
     "used throughout the UAMS React frontend."),
    ("UAMS",
     "University Asset Management System — the name of the project described in this thesis."),
    ("VK",
     "VKontakte — a Russian social network whose messaging API is integrated into UAMS "
     "to allow reporters to submit asset defect reports via a VK bot."),
]

for term, definition in glossary:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0); pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0)
    run1 = p.add_run(f"{term} — ")
    run1.bold = True
    run1.font.name = "Times New Roman"; run1.font.size = Pt(12)
    run2 = p.add_run(definition)
    run2.font.name = "Times New Roman"; run2.font.size = Pt(12)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════
add_heading("INTRODUCTION", level=1)

body(
    "Every university campus contains thousands of physical assets: chairs, desks, projectors, "
    "computers, whiteboards, and the infrastructure rooms that house them. Keeping an accurate, "
    "up-to-date record of where each asset is located, what condition it is in, and which "
    "maintenance tasks have been completed is a non-trivial operational challenge. In many "
    "institutions this work is still handled through paper ledgers, scattered spreadsheets, or "
    "ad-hoc email threads — approaches that do not scale and are prone to human error."
)
body(
    "The University Asset Management System (UAMS) was conceived to address this gap. It "
    "provides a web-based platform that gives asset managers a visual, interactive canvas on "
    "which to design room layouts, place named asset instances, and track their state over time. "
    "Reporters (e.g., lecturers or lab technicians) can submit defect notifications directly "
    "through a VK social-media bot, which routes the report to the responsible manager without "
    "requiring them to use the web interface at all."
)
body(
    "The relevance of this work lies in the intersection of two active areas of software "
    "engineering: domain-driven design for enterprise applications and rich-client browser "
    "interfaces that replicate desktop interaction patterns. Building a drag-and-drop canvas "
    "without a dedicated graphics library, while maintaining synchronisation with a strongly "
    "typed backend, requires careful architectural decisions that this thesis documents and "
    "analyses."
)

add_heading("Objective", level=2)
body(
    "The objective of this work is to design and implement a production-ready web application "
    "for managing university physical assets, covering the full stack from database schema to "
    "interactive browser canvas."
)

add_heading("Tasks", level=2)
tasks = [
    "Analyse the functional requirements of university asset management.",
    "Design the problem domain model and database schema.",
    "Select and justify the technology stack for frontend and backend.",
    "Implement the backend REST API following Clean Architecture and CQRS principles.",
    "Implement the frontend SPA including the interactive canvas designer.",
    "Integrate a VK bot for reporter-facing defect submission.",
    "Evaluate the implemented system against the stated requirements.",
]
for i, t in enumerate(tasks, 1):
    p = doc.add_paragraph(f"{i}. {t}")
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0); pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25)
    for run in p.runs:
        run.font.name = "Times New Roman"; run.font.size = Pt(12)

add_heading("Structure of the Work", level=2)
body(
    "This thesis consists of four chapters. Chapter 1 covers the design phase: stakeholder "
    "analysis, functional requirements, and the domain model. Chapter 2 justifies the selected "
    "technologies. Chapter 3 describes the overall system architecture. Chapter 4 documents "
    "key implementation details for each major feature area. The thesis concludes with a "
    "summary of results and directions for future development."
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# CHAPTER 1 — DESIGN
# ═══════════════════════════════════════════════════════════════════════════
add_heading("1  DESIGN", level=1)

add_heading("1.1  Functional Requirements", level=2)

body(
    "The requirements were derived through iterative discussions with the intended users of the "
    "system: faculty asset managers responsible for maintaining inventories, and reporters who "
    "interact with assets day-to-day. The resulting functional requirements are grouped by role."
)

add_heading("Asset Manager Requirements", level=3)
am_reqs = [
    "Create, edit, and deactivate Asset Definitions with a name, category, and SVG icon.",
    "Create Buildings and Rooms, assigning each Room to a Faculty and a Building.",
    "Open a Room in the Canvas Designer and place, move, resize, rotate, and delete "
    "asset instances (Placed Assets) on the canvas.",
    "Group multiple Placed Assets under a shared label so they can be selected and moved together.",
    "Link non-infrastructure assets to an infrastructure (room) asset on the canvas so that "
    "moving the room moves all linked children.",
    "Change the condition of any Placed Asset (Good, Reported, Under Maintenance, "
    "Irreparable, Replaced).",
    "View and complete per-asset maintenance checklists, which are auto-generated from the "
    "Asset Definition's checklist template at the start of each study year.",
    "Close or reopen a Room, recording a reason for closure.",
]
for i, r in enumerate(am_reqs, 1):
    p = doc.add_paragraph(f"AM-{i:02d}. {r}")
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0); pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25)
    for run in p.runs:
        run.font.name = "Times New Roman"; run.font.size = Pt(12)

add_heading("Reporter Requirements", level=3)
rep_reqs = [
    "Browse the list of buildings and rooms accessible to them.",
    "View the read-only canvas layout of any open room.",
    "Submit a defect report for a specific Placed Asset via the web interface or via VK bot.",
    "Receive confirmation when a defect report is accepted.",
]
for i, r in enumerate(rep_reqs, 1):
    p = doc.add_paragraph(f"RP-{i:02d}. {r}")
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0); pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25)
    for run in p.runs:
        run.font.name = "Times New Roman"; run.font.size = Pt(12)

add_heading("Non-Functional Requirements", level=3)
nfr = [
    "Security: all write operations require a valid JWT and a role claim check.",
    "Consistency: the canvas layout is stored atomically as a JSON snapshot; partial saves "
    "are not possible.",
    "Maintainability: the backend must follow a layered architecture with no cross-layer "
    "dependency violations.",
    "Usability: the canvas must support zoom and pan without performance degradation for "
    "rooms with up to 200 assets.",
]
for i, r in enumerate(nfr, 1):
    p = doc.add_paragraph(f"NFR-{i:02d}. {r}")
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0); pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25)
    for run in p.runs:
        run.font.name = "Times New Roman"; run.font.size = Pt(12)

add_heading("1.2  Problem Domain Model", level=2)

body(
    "Figure 1.1 shows the simplified class diagram of the UAMS domain. The top-level "
    "aggregate is Faculty, which owns a collection of Buildings. Each Building contains "
    "Rooms. A Room has exactly one Layout, which in turn owns a collection of Placed Assets. "
    "Each Placed Asset references one Asset Definition."
)
body(
    "The Faculty entity records the Id of the designated Asset Manager and acts as the "
    "authorization boundary: only the asset manager of a given Faculty may modify the rooms "
    "belonging to that Faculty."
)
body(
    "An Asset Definition captures the reusable metadata for a class of asset: its display "
    "name, category (Infrastructure, Furniture, Electronics, Other), the URL of its SVG icon, "
    "and an ordered list of checklist items. Multiple Placed Asset instances on different room "
    "canvases can share the same Asset Definition."
)
body(
    "A Placed Asset is the concrete, positioned instance of a definition inside a specific "
    "room. It records the asset's pixel-space coordinates (X, Y, Width, Height, Rotation), "
    "its current Condition, and two optional grouping/linking fields: GroupId/GroupLabel for "
    "multi-selection groups and CanvasRoomId for linking a non-infrastructure asset to an "
    "infrastructure (room) asset on the canvas."
)
body(
    "A Placed Asset Checklist is created automatically when a Placed Asset is first placed "
    "and carries a study year label (e.g., '2025–2026'). It contains a list of Checklist "
    "Items, each with a description, a completion flag, and an optional note."
)
body(
    "[Figure 1.1 – Domain class diagram. See Appendix A.]"
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# CHAPTER 2 — TECHNOLOGIES
# ═══════════════════════════════════════════════════════════════════════════
add_heading("2  TECHNOLOGIES AND TOOLS", level=1)

add_heading("2.1  React, TypeScript and Vite", level=2)
body(
    "The frontend is a React 18 single-page application written entirely in TypeScript 5. "
    "React was chosen because its component model and unidirectional data flow suit the "
    "canvas-heavy interface well: each asset tile, handle, and overlay is a discrete component "
    "whose re-rendering can be controlled precisely through memo and callback stabilisation."
)
body(
    "TypeScript was adopted to eliminate an entire class of runtime errors — incorrect property "
    "names, missing fields, type mismatches between API responses and local state — that would "
    "be difficult to catch through testing alone in a large canvas interaction loop. Strict "
    "mode is enabled throughout the project."
)
body(
    "Vite serves as the build tool and development server. Its ES-module-native architecture "
    "produces sub-second hot-module replacement (HMR) cycles even for large files such as "
    "CanvasPage.tsx, which contains the entire canvas interaction state machine. "
    "Tailwind CSS is used for styling; its utility-class approach avoids style name collisions "
    "and keeps component markup self-documenting."
)

add_heading("2.2  ASP.NET Core", level=2)
body(
    "The backend is an ASP.NET Core 8 Web API project. ASP.NET Core was selected over "
    "alternatives such as Node.js/Express or Django because the team had prior experience with "
    "the .NET ecosystem, the framework provides first-class support for the patterns used "
    "(dependency injection, middleware, controllers), and its performance profile is well-suited "
    "to I/O-bound workloads like database access."
)
body(
    "The API project exposes two areas of controllers: one for room and asset management "
    "and one for authentication delegation. All business logic is kept out of controllers; "
    "they act solely as thin MediatR dispatchers."
)

add_heading("2.2.1  Entity Framework Core", level=3)
body(
    "Entity Framework Core 8 (EF Core) is the data access layer. The schema is maintained "
    "code-first: migrations are generated from the domain model classes and applied to a "
    "SQL Server database. A key architectural decision was to store the entire canvas Layout "
    "— including all Placed Assets — as a JSON-owned column on the Room entity using EF Core's "
    "OwnsOne(...).ToJson() API."
)
body(
    "This approach was chosen because the Layout is always read and written as a single atomic "
    "unit (a full snapshot is saved each time the asset manager clicks Save). Normalising each "
    "Placed Asset into its own row would introduce unnecessary join complexity and would provide "
    "no benefit for the access pattern actually used. The JSON column also means that adding "
    "new optional fields to Placed Assets (such as CanvasRoomId) does not require database "
    "migrations — the column simply stores wider JSON objects going forward."
)

add_heading("2.2.2  MediatR and CQRS", level=3)
body(
    "MediatR is used to implement the CQRS pattern. Every feature is expressed as either a "
    "Command (mutating operation) or a Query (read operation), each paired with a dedicated "
    "handler class. This structure enforces a clean separation of concerns: the handler for "
    "SaveLayoutCommand contains all the logic for validating and persisting a canvas snapshot, "
    "completely independently of the HTTP layer."
)
body(
    "The pattern also simplifies testing: a command handler can be unit-tested by constructing "
    "it directly with a mock DbContext, without spinning up an HTTP server. Pipeline behaviours "
    "in MediatR are used to attach cross-cutting concerns — validation and logging — without "
    "modifying handler code."
)

add_heading("2.3  SQL Server", level=2)
body(
    "Microsoft SQL Server 2022 is used as the relational database. The choice was driven by "
    "compatibility with EF Core's JSON column support (introduced fully in EF Core 8 for SQL "
    "Server), as well as the operational familiarity of the deployment environment. All "
    "identifiers are GUIDs generated client-side before the HTTP request is sent; this avoids "
    "round-trips to retrieve auto-increment keys and allows the frontend to reference newly "
    "created assets immediately."
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# CHAPTER 3 — ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════
add_heading("3  ARCHITECTURE", level=1)

body(
    "The UAMS backend follows the Clean Architecture pattern, organised into four layers that "
    "enforce a strict inward-only dependency rule: outer layers may depend on inner layers, "
    "but inner layers know nothing about outer layers."
)

add_heading("3.1  API Layer", level=2)
body(
    "The outermost layer consists of ASP.NET Core controllers, middleware, and dependency "
    "injection registrations. Controllers receive HTTP requests, map them to MediatR commands "
    "or queries, dispatch them, and return HTTP responses. No business logic resides here. "
    "Route parameters and request bodies are mapped to strongly-typed records — the same "
    "records that form the command/query API surface exposed to the Application layer."
)
body(
    "Authentication is handled by a JWT validation middleware that reads the Bearer token, "
    "validates its signature, and populates the HttpContext user identity with role claims. "
    "A custom IFacultyFacade interface provides the Application layer with a way to check "
    "role ownership without coupling it to the HTTP context directly."
)

add_heading("3.2  Application Layer", level=2)
body(
    "The Application layer contains all CQRS command and query handlers plus the DTO and "
    "view record definitions returned to callers. This layer depends only on the Domain layer "
    "and on abstractions (interfaces) that are implemented by outer layers."
)
body(
    "Each feature folder (e.g., SaveLayoutFeature, GetRoomByIdFeature) is self-contained: "
    "it holds the request record, the handler, and any local validation logic. This vertical "
    "slice organisation keeps related code co-located and minimises the cost of adding or "
    "removing features."
)

add_heading("3.3  Persistence Layer", level=2)
body(
    "The Persistence layer provides the EF Core DbContext (RoomDesignDbContext) and the "
    "Fluent API configuration for entity mappings. The Room → Layout → PlacedAssets "
    "hierarchy is configured with OwnsOne and OwnsMany combined with ToJson(), meaning that "
    "EF Core serialises and deserialises the entire layout as a single JSON value without the "
    "application code needing to do any manual JSON handling."
)

add_heading("3.4  Domain Layer", level=2)
body(
    "The Domain layer contains plain C# entity classes, value objects, and domain-specific "
    "logic. The most significant domain method is Layout.ApplySnapshot(), which accepts the "
    "incoming list of Placed Assets and computes a diff (added IDs, removed IDs) against the "
    "current state. The diff is returned to the handler so it can trigger side-effects "
    "(creating or deleting checklists) without the domain object knowing about the database."
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# CHAPTER 4 — IMPLEMENTATION
# ═══════════════════════════════════════════════════════════════════════════
add_heading("4  IMPLEMENTATION", level=1)

add_heading("4.1  Authentication and Role-Based Access Control", level=2)
body(
    "UAMS uses JWT-based authentication. Tokens are issued by an external identity endpoint "
    "and carry two relevant claims: the user's GUID (sub) and their role (AssetManager or "
    "Reporter). The backend does not issue tokens; it only validates them."
)
body(
    "Role enforcement happens in two places. First, ASP.NET Core's [Authorize(Roles=...)] "
    "attribute on controllers provides coarse-grained protection at the route level. Second, "
    "within command handlers, the IFacultyFacade.IsAssetManagerOfFaculty(userId, facultyId) "
    "call provides fine-grained ownership checks — confirming that the authenticated user is "
    "the manager of the specific faculty containing the room being modified, not merely any "
    "asset manager in the system."
)
body(
    "On the frontend, the JWT payload is decoded on login, and the role claim is stored in "
    "React context. Route guards implemented as higher-order components redirect users to the "
    "appropriate dashboard based on their role, and UI elements that trigger write operations "
    "are conditionally rendered only for asset managers."
)

add_heading("4.2  Asset Definitions", level=2)
body(
    "Asset Definitions are the catalogue entries from which asset instances on canvases are "
    "created. Each definition records a unique GUID, a display name, a category enum "
    "(Infrastructure, Furniture, Electronics, Other), the URL of an SVG icon hosted externally "
    "(GitHub raw content), and an ordered list of checklist item descriptions."
)
body(
    "The CreateAssetDefinitionCommand validates that the name is non-empty and that the "
    "category is a recognised enum value before persisting. The SVG URL is stored verbatim; "
    "the frontend fetches and renders it at display time. Because the icons are SVG, they "
    "scale without quality loss regardless of the zoom level or the pixel dimensions assigned "
    "to the asset on the canvas."
)

add_heading("4.3  Buildings and Rooms", level=2)
body(
    "The building/room hierarchy reflects the physical structure of a campus. A Building "
    "belongs to a Faculty and holds Rooms. A Room records its name, the Faculty it belongs "
    "to, a status (Open or Closed), an optional closure reason, and the GUID of the user who "
    "last designed its layout."
)
body(
    "When a Room is created, EF Core automatically initialises its owned Layout with an empty "
    "PlacedAssets list via the OwnsOne configuration. This means the canvas is always "
    "non-null on the server side; the frontend never needs to handle a missing layout."
)

add_heading("4.4  Canvas Designer", level=2)
body(
    "The Canvas Designer is the most complex component in the frontend. It is implemented as "
    "a single large functional component (CanvasPage.tsx) that maintains the full interaction "
    "state in React hooks and translates pointer events into asset transformations."
)
body(
    "The canvas is rendered as a div with CSS transform: scale(zoom) translate(panX, panY). "
    "Each Placed Asset is an absolutely-positioned div overlaid with eight resize handles "
    "(corner and edge). This approach avoids a WebGL or Canvas 2D API, keeping the rendering "
    "fully in the browser's normal layout engine and making it trivial to embed React "
    "components (tooltips, overlays) directly on top of assets."
)

add_heading("4.4.1  Placing, Moving and Resizing", level=3)
body(
    "Placing an asset creates a new entry in the local assets state array with a "
    "freshly-generated client-side GUID. Moving is implemented via an interactionRef that "
    "stores the type of interaction (drag, or one of eight resize directions) and the pointer "
    "delta since the last frame."
)
body(
    "Resizing maintains aspect ratio when a corner handle is dragged while the Shift key is "
    "held. The resize handle size is intentionally zoom-invariant: handle size = 14 / zoom "
    "pixels and offset = −7 / zoom pixels. This ensures that handles remain easily "
    "clickable regardless of the current zoom level — if the handles were sized in canvas "
    "coordinates they would become too small to click when zoomed out."
)
body(
    "SVG icons are rendered as HTML <img> elements with the object-fit: fill CSS property. "
    "Each SVG file carries the preserveAspectRatio=\"none\" attribute, which instructs the "
    "browser's SVG renderer to stretch the graphic to fill its viewport exactly, matching "
    "the behaviour users expect when they resize an asset."
)

add_heading("4.4.2  Grouping and Room Linking", level=3)
body(
    "The grouping feature allows the asset manager to select multiple assets and assign them "
    "a shared GroupId and GroupLabel. A grouped drag operation moves all members of the group "
    "together by iterating over the selected subset and applying the same delta to each."
)
body(
    "Room linking (CanvasRoomId) solves the problem of moving an infrastructure room asset "
    "along with the furniture it contains. Each non-infrastructure Placed Asset records the "
    "optional GUID of the infrastructure asset it sits inside. This value is computed "
    "automatically on every pointer-up event: the handler iterates all non-infrastructure "
    "assets, computes their centre point, and finds the infrastructure asset whose bounding "
    "box contains that point. If the centre is inside a room, CanvasRoomId is set to that "
    "room's ID; if it is outside all rooms, CanvasRoomId is set to null."
)
body(
    "On the backend, CanvasRoomId is persisted as part of the JSON layout snapshot. No "
    "database migration was required because the JSON column accepts new fields transparently; "
    "legacy rows simply do not have the key and EF Core deserialises it as null."
)

add_heading("4.5  Asset Conditions and Checklists", level=2)
body(
    "Every Placed Asset carries a Condition enum with five values: Good, Reported, "
    "UnderMaintenance, Irreparable, and Replaced. The condition is visualised on the canvas "
    "as a coloured border: transparent for Good, amber for Reported, blue for Under "
    "Maintenance, red for Irreparable, and teal for Replaced. This gives the asset manager "
    "an at-a-glance health overview of an entire room without opening individual records."
)
body(
    "Checklists are created automatically by the SaveLayoutCommand handler when a new Placed "
    "Asset is added to the canvas (detected by comparing the incoming snapshot against the "
    "current one). The study year label is computed server-side: September onwards belongs to "
    "the new academic year (e.g., September 2025 → '2025–2026')."
)
body(
    "Each checklist item records a description (copied from the Asset Definition template), "
    "a boolean IsCompleted flag, and an optional Note field for the manager to record "
    "observations. Completing checklist items does not automatically change the asset's "
    "Condition; the manager sets the condition explicitly."
)

add_heading("4.6  Reporting and VK Bot Integration", level=2)
body(
    "The Reporter role can submit defect reports through two channels: the web interface and "
    "a VK social-media bot. Both channels produce the same backend command "
    "(SubmitAssetReportCommand) which sets the Placed Asset's condition to Reported and "
    "records the report details."
)
body(
    "The VK bot is implemented as a long-polling service that connects to the VK API and "
    "listens for incoming messages. When a reporter sends a message to the bot, the bot "
    "guides them through a multi-step conversation: selecting the building, selecting the "
    "room, selecting the asset from the room's layout, and describing the defect. "
    "The bot then calls the UAMS REST API internally to record the report."
)
body(
    "This design decouples the bot from the web frontend entirely; the bot is a separate "
    "hosted service that speaks the same HTTP API as the browser. Reporters who are "
    "uncomfortable with web applications can therefore participate in the reporting workflow "
    "through a familiar chat interface."
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════
add_heading("CONCLUSION", level=1)

body(
    "This thesis has described the design and implementation of the University Asset "
    "Management System (UAMS), a full-stack web application that replaces manual asset "
    "tracking processes with an interactive, role-aware digital platform."
)
body(
    "The main results of the work are as follows. A domain model was designed that captures "
    "the Faculty → Building → Room → Layout → Placed Asset hierarchy and the relationships "
    "between asset definitions, placed instances, conditions, and checklists. A Clean "
    "Architecture backend was implemented in ASP.NET Core with EF Core and MediatR, "
    "demonstrating that the CQRS pattern leads to handler code that is concise, testable, "
    "and easy to extend. A React/TypeScript SPA was implemented featuring a custom canvas "
    "designer with zoom, pan, drag, resize, grouping, and room-linking capabilities — "
    "all without a dedicated graphics library. A VK bot was integrated to provide a "
    "chat-native reporting channel."
)
body(
    "The most technically challenging aspect of the project was the canvas interaction system. "
    "Managing pointer events, zoom-invariant handle sizing, SVG scaling, and bidirectional "
    "room linking in a single React component required careful use of refs to avoid stale "
    "closures and functional state updates to ensure correct recalculation on every drag end."
)
body(
    "Directions for future development include: replacing the manual canvas renderer with a "
    "canvas-native library (e.g., Konva.js) for better performance at very large asset counts; "
    "adding real-time collaboration via WebSockets so multiple managers can edit a room "
    "simultaneously; extending the reporting flow to include photo uploads; and implementing "
    "analytics dashboards that aggregate condition data across buildings."
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════════════
add_heading("REFERENCES", level=1)

refs = [
    "Martin, R. C. Clean Architecture: A Craftsman's Guide to Software Structure and Design. "
    "— Prentice Hall, 2017. — 432 p.",

    "Fowler, M. Patterns of Enterprise Application Architecture. "
    "— Addison-Wesley, 2002. — 560 p.",

    "Microsoft. ASP.NET Core documentation [Electronic resource]. "
    "— URL: https://learn.microsoft.com/aspnet/core (accessed 20.05.2026).",

    "Microsoft. Entity Framework Core documentation [Electronic resource]. "
    "— URL: https://learn.microsoft.com/ef/core (accessed 20.05.2026).",

    "Young, G. CQRS Documents [Electronic resource]. "
    "— URL: https://cqrs.files.wordpress.com/2010/11/cqrs_documents.pdf (accessed 15.04.2026).",

    "Richardson, C. Microservices Patterns. — Manning, 2018. — 520 p.",

    "React. React Documentation [Electronic resource]. "
    "— URL: https://react.dev (accessed 20.05.2026).",

    "TypeScript. TypeScript Handbook [Electronic resource]. "
    "— URL: https://www.typescriptlang.org/docs/handbook (accessed 20.05.2026).",

    "Tailwind CSS. Tailwind CSS Documentation [Electronic resource]. "
    "— URL: https://tailwindcss.com/docs (accessed 20.05.2026).",

    "VK. VK API Documentation [Electronic resource]. "
    "— URL: https://dev.vk.com/reference (accessed 20.05.2026).",

    "W3C. Scalable Vector Graphics (SVG) 1.1 Specification [Electronic resource]. "
    "— URL: https://www.w3.org/TR/SVG11 (accessed 10.03.2026).",

    "Newman, S. Building Microservices: Designing Fine-Grained Systems, 2nd ed. "
    "— O'Reilly, 2021. — 616 p.",
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph(f"{i}. {ref}")
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0); pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0)
    for run in p.runs:
        run.font.name = "Times New Roman"; run.font.size = Pt(12)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# APPENDIX A
# ═══════════════════════════════════════════════════════════════════════════
add_heading("APPENDIX A – SELECTED SCREENSHOTS", level=1)

body(
    "The following figures illustrate the main screens of the UAMS application. "
    "Screenshots should be inserted here by the student before final submission."
)

screenshots = [
    ("Figure A.1", "Asset Manager Dashboard — building and room list."),
    ("Figure A.2", "Canvas Designer — room layout with placed assets, groups, and "
                   "condition colour coding."),
    ("Figure A.3", "Canvas Designer — resize handles visible on a selected asset."),
    ("Figure A.4", "Asset Definition management page — list and creation form."),
    ("Figure A.5", "Placed Asset checklist — checklist items for a specific asset."),
    ("Figure A.6", "VK bot conversation — reporter submitting a defect report."),
]

for fig, caption in screenshots:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(6); pf.space_after = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0)
    run1 = p.add_run(f"{fig} — ")
    run1.bold = True
    run1.font.name = "Times New Roman"; run1.font.size = Pt(12)
    run2 = p.add_run(caption)
    run2.font.name = "Times New Roman"; run2.font.size = Pt(12)

    # Placeholder box
    p2 = doc.add_paragraph("[Insert screenshot here]")
    pf2 = p2.paragraph_format
    pf2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf2.space_before = Pt(2); pf2.space_after = Pt(12)
    pf2.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf2.line_spacing = 1.5
    pf2.first_line_indent = Cm(0)
    for run in p2.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.italic = True

# ─── Save ──────────────────────────────────────────────────────────────────
out_path = r"c:\Users\user\Desktop\UAMS_Thesis.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
